import os, io, json, uuid, hashlib, tempfile, subprocess
from pathlib import Path
from typing import List, Dict, Optional
import fitz  # PyMuPDF
from pydantic import BaseModel
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse
import gradio as gr
import uvicorn
import requests
from qdrant_client import QdrantClient
from qdrant_client.http.models import Distance, VectorParams
from datetime import datetime
import asyncio
import threading
from concurrent.futures import ThreadPoolExecutor

APP_PORT = int(os.getenv("APP_PORT", "7860"))
QDRANT_URL = os.getenv("QDRANT_URL", "http://localhost:6333")
COLLECTION = "contract_chunks"
HUB_GATEWAY_URL = os.getenv("HUB_GATEWAY_URL", "http://hub-gateway:8081")
POLICY_FILE = os.getenv("POLICY_FILE", "/app/policy/nda.yml")
DATA_DIR = Path(os.getenv("ABS_DATA_DIR", "/data"))
REPORTS_DIR = DATA_DIR / "reports"
UPLOADS_DIR = DATA_DIR / "uploads"
DATA_DIR.mkdir(parents=True, exist_ok=True)
REPORTS_DIR.mkdir(parents=True, exist_ok=True)
UPLOADS_DIR.mkdir(parents=True, exist_ok=True)

# Global state for review history
review_history = []
executor = ThreadPoolExecutor(max_workers=4)

from docx import Document as DocxDocument

ALLOWED_EXTS = {".pdf", ".docx"}

def read_docx_text(path: Path) -> str:
    doc = DocxDocument(str(path))
    return "\n".join(p.text for p in doc.paragraphs if p.text)

def read_any_text(path: Path) -> tuple[str, str]:
    ext = path.suffix.lower()
    if ext not in ALLOWED_EXTS:
        raise HTTPException(status_code=415, detail=f"Unsupported file type: {ext}. Please upload PDF or DOCX.")
    if ext == ".pdf":
        pdf_path = ocr_if_scanned(path)
        return read_pdf_text(pdf_path), "pdf"
    return read_docx_text(path), "docx"



# ----- models/clients -----
qdrant = QdrantClient(url=QDRANT_URL)

# Initialize collection with legal embedding dimension
EMBED_DIM = 768  # Legal-BERT dimension
try:
    qdrant.get_collection(COLLECTION)
except Exception:
    qdrant.recreate_collection(
        COLLECTION,
        vectors_config=VectorParams(size=EMBED_DIM, distance=Distance.COSINE)
    )

def sha256_of_bytes(b: bytes) -> str:
    return hashlib.sha256(b).hexdigest()

def ocr_if_scanned(pdf_path: Path) -> Path:
    out = pdf_path.with_suffix(".ocr.pdf")
    try:
        subprocess.run(["ocrmypdf", "--skip-text", str(pdf_path), str(out)], check=True)
        return out if out.exists() else pdf_path
    except Exception:
        return pdf_path

def read_pdf_text(path: Path) -> str:
    doc = fitz.open(str(path))
    texts = []
    for page in doc:
        texts.append(page.get_text())
    return "\n".join(texts)

def chunk_text(text: str, max_chars=1200, overlap=120):
    chunks = []
    i=0
    while i < len(text):
        chunk = text[i:i+max_chars]
        chunks.append(chunk)
        i += max_chars - overlap
    return chunks

def upsert_chunks(chunks: List[str], doc_id: str):
    # Get embeddings from Hub Gateway
    response = requests.post(
        f"{HUB_GATEWAY_URL}/v1/embeddings",
        headers={"X-ABS-App-Id": "contract-reviewer"},
        json={"input": chunks},
        timeout=60
    )
    response.raise_for_status()
    embeddings_data = response.json()
    
    vectors = [item["embedding"] for item in embeddings_data["data"]]
    ids = [str(uuid.uuid4()) for _ in chunks]
    qdrant.upsert(
        COLLECTION,
        points=[{"id": ids[i], "vector": vectors[i], "payload": {"text": chunks[i], "doc_id": doc_id}} for i in range(len(chunks))]
    )

def retrieve(query: str, top_k=6, doc_id: Optional[str] = None) -> List[str]:
    # Get query embedding from Hub Gateway
    response = requests.post(
        f"{HUB_GATEWAY_URL}/v1/embeddings",
        headers={"X-ABS-App-Id": "contract-reviewer"},
        json={"input": [query]},
        timeout=60
    )
    response.raise_for_status()
    embeddings_data = response.json()
    
    qvec = embeddings_data["data"][0]["embedding"]
    search_filter = None
    if doc_id:
        from qdrant_client.http import models as qmodels
        search_filter = qmodels.Filter(must=[qmodels.FieldCondition(key="doc_id", match=qmodels.MatchValue(value=doc_id))])
    res = qdrant.search(COLLECTION, query_vector=qvec, limit=top_k, query_filter=search_filter)
    return [hit.payload["text"] for hit in res]

def openai_chat(messages, max_tokens=1024) -> str:
    # Use Hub Gateway for chat completions
    response = requests.post(
        f"{HUB_GATEWAY_URL}/v1/chat/completions",
        headers={"X-ABS-App-Id": "contract-reviewer"},
        json={
        "messages": messages,
            "temperature": 0.2,
            "max_tokens": max_tokens
        },
        timeout=120
    )
    response.raise_for_status()
    result = response.json()
    return result["choices"][0]["message"]["content"]

def review_contract(text: str, policy: str, doc_id: Optional[str]) -> Dict:
    # collect context via retrieval queries for common clauses
    queries = [
        "Confidentiality clause", "Liability limitation", "Indemnity",
        "Intellectual Property ownership", "Termination", "Governing Law",
        "Assignment", "Data Protection / Privacy"
    ]
    contexts = []
    for q in queries:
        contexts.extend(retrieve(q, doc_id=doc_id))
    context = "\n---\n".join(contexts[:12])

    with open("/app/prompts/review_prompt.md", "r", encoding="utf-8") as f:
        system_prompt = f.read()

    # Include a substantial document excerpt to anchor the model on the uploaded file
    doc_excerpt = text[:6000]  # Increased excerpt size
    user_prompt = f"""DOCUMENT TO ANALYZE:\n{doc_excerpt}\n\nPOLICY GUIDELINES:\n{policy}\n\nRELEVANT CONTEXT FROM SIMILAR DOCUMENTS:\n{context}\n\nINSTRUCTIONS: Analyze the DOCUMENT TO ANALYZE above and return ONLY valid JSON as specified in the system prompt."""
    rsp = openai_chat([
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt}
    ], max_tokens=int(os.getenv("MAX_TOKENS", "1024")))

    # Debug: Log the raw response
    print(f"DEBUG: Raw LLM response: {repr(rsp)}")
    
    try:
        # Clean the response - remove markdown code blocks and extra text
        cleaned_rsp = rsp.strip()
        
        # Remove markdown code blocks if present
        if cleaned_rsp.startswith('```'):
            lines = cleaned_rsp.split('\n')
            if len(lines) > 2:
                cleaned_rsp = '\n'.join(lines[1:-1])
        
        # Remove explanatory text before JSON
        json_start = cleaned_rsp.find('{')
        if json_start != -1:
            cleaned_rsp = cleaned_rsp[json_start:]
        
        # Find the end of JSON object
        json_end = cleaned_rsp.rfind('}') + 1
        if json_end > 0:
            cleaned_rsp = cleaned_rsp[:json_end]
        
        print(f"DEBUG: Cleaned JSON string: {repr(cleaned_rsp)}")
        data = json.loads(cleaned_rsp)
        
        # Fallback: If LLM returns wrong structure, create expected structure
        if "summary" not in data:
            print("DEBUG: Wrong JSON structure detected, creating fallback")
            
            # Create a meaningful summary from the document content
            summary_parts = []
            risk_count = 0
            key_terms = []
            
            # Extract meaningful information from wrong structure
            if "agreement" in data:
                agreement = data["agreement"]
                if "terms" in agreement:
                    key_terms = [term.get('text', '') for term in agreement["terms"][:3] if term.get('text')]
                
                # Count risks from various sources
                if "representations" in agreement:
                    risk_count += len(agreement["representations"])
                if "liabilities" in agreement:
                    risk_count += len(agreement["liabilities"])
                if "disclaimers" in agreement:
                    risk_count += len(agreement["disclaimers"])
            
            if "prohibitions" in data:
                risk_count += len(data["prohibitions"])
            if "liquidated_damages" in data:
                risk_count += len(data["liquidated_damages"])
            
            # Create meaningful summary
            if key_terms:
                summary_parts.append(f"This document appears to be a legal agreement covering: {', '.join(key_terms[:3])}.")
            else:
                summary_parts.append("This document appears to be a legal agreement or contract.")
            
            if risk_count > 0:
                summary_parts.append(f"The document contains {risk_count} potential risk areas that require careful review.")
            else:
                summary_parts.append("The document appears to be relatively standard with no obvious high-risk provisions.")
            
            summary_parts.append("Please review the detailed analysis below for specific recommendations.")
            
            fallback_data = {
                "summary": " ".join(summary_parts),
                "document_type": "Contract",
                "key_points": [],
                "risks": [],
                "recommendations": [],
                "citations": []
            }
            
            # Try to extract useful info from wrong structure
            if "agreement" in data:
                agreement = data["agreement"]
                if "terms" in agreement:
                    # Create meaningful key points from terms
                    key_points = []
                    for term in agreement["terms"][:5]:
                        section = term.get('section', '')
                        text = term.get('text', '')
                        if section and text and text != section:
                            key_points.append(f"{section}: {text}")
                        elif section:
                            key_points.append(section)
                    fallback_data["key_points"] = key_points
                if "representations" in agreement:
                    fallback_data["risks"].append({
                        "level": "Medium",
                        "description": "Representations and warranties identified",
                        "rationale": "Contract contains specific representations that may create liability"
                    })
                if "liabilities" in agreement:
                    for liability in agreement["liabilities"]:
                        fallback_data["risks"].append({
                            "level": "High",
                            "description": liability.get("description", "Liability clause identified"),
                            "rationale": "Contract contains liability provisions"
                        })
            
            if "prohibitions" in data:
                for prohibition in data["prohibitions"]:
                    fallback_data["risks"].append({
                        "level": "Medium",
                        "description": prohibition.get("description", "Restriction identified"),
                        "rationale": "Contract contains restrictive provisions"
                    })
            
            if "liquidated_damages" in data:
                for damage in data["liquidated_damages"]:
                    fallback_data["risks"].append({
                        "level": "High",
                        "description": f"Liquidated damages: {damage.get('amount', 'Amount specified')}",
                        "rationale": "Contract contains liquidated damages clause"
                    })
            
            # Add generic recommendations if none found
            if not fallback_data["recommendations"]:
                fallback_data["recommendations"] = [
                    "Review all liability and damage clauses carefully",
                    "Consider legal counsel for complex provisions",
                    "Verify compliance with applicable laws"
                ]
            
            data = fallback_data
        
    except json.JSONDecodeError as e:
        # If JSON parsing fails, wrap the response
        print(f"DEBUG: JSON parsing failed: {str(e)}")
        data = {"summary": rsp, "raw": rsp, "error": f"JSON parsing failed: {str(e)}"}
    except Exception as e:
        # Any other error
        print(f"DEBUG: Processing failed: {str(e)}")
        data = {"summary": rsp, "raw": rsp, "error": f"Processing failed: {str(e)}"}

    return data

# -------- FastAPI + Gradio --------
app = FastAPI()

class ReviewResponse(BaseModel):
    summary: str
    clauses: list | None = None
    risks: list | None = None
    missing: list | None = None
    citations: list | None = None
    input_sha256: str

@app.get("/healthz")
def healthz():
    return {"status": "ok"}

@app.get("/api/history")
def get_history():
    """Get review history"""
    return {"history": review_history[-50:]}  # Last 50 reviews

@app.get("/api/reports/{report_id}")
def get_report(report_id: str):
    """Get a specific report by ID"""
    report_file = REPORTS_DIR / f"report-{report_id}.json"
    if not report_file.exists():
        raise HTTPException(status_code=404, detail="Report not found")
    
    with open(report_file, "r", encoding="utf-8") as f:
        return json.load(f)

@app.delete("/api/reports/{report_id}")
def delete_report(report_id: str):
    """Delete a specific report"""
    report_file = REPORTS_DIR / f"report-{report_id}.json"
    if not report_file.exists():
        raise HTTPException(status_code=404, detail="Report not found")
    
    report_file.unlink()
    
    # Remove from history
    global review_history
    review_history = [h for h in review_history if h["id"] != report_id]
    
    return {"message": "Report deleted successfully"}

@app.get("/api/export/{report_id}/pdf")
def export_pdf(report_id: str):
    """Export report as PDF"""
    report_file = REPORTS_DIR / f"report-{report_id}.json"
    if not report_file.exists():
        raise HTTPException(status_code=404, detail="Report not found")
    
    with open(report_file, "r", encoding="utf-8") as f:
        report_data = json.load(f)
    
    # Generate PDF content
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    
    # Create PDF buffer
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
    
    # Get styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=18, spaceAfter=30, alignment=TA_CENTER)
    heading_style = ParagraphStyle('CustomHeading', parent=styles['Heading2'], fontSize=14, spaceAfter=12)
    normal_style = styles['Normal']
    
    # Build PDF content
    story = []
    
    # Title
    story.append(Paragraph("Contract Review Report", title_style))
    story.append(Spacer(1, 12))
    
    # Document info
    story.append(Paragraph(f"<b>Document:</b> {report_data.get('file_name', 'Unknown')}", normal_style))
    story.append(Paragraph(f"<b>Analysis Date:</b> {report_data.get('timestamp', 'Unknown')}", normal_style))
    story.append(Spacer(1, 20))
    
    # Summary
    story.append(Paragraph("Executive Summary", heading_style))
    summary_text = report_data.get('summary', 'No summary available')
    story.append(Paragraph(summary_text, normal_style))
    story.append(Spacer(1, 20))
    
    # Risk Analysis
    risks = report_data.get('risks', [])
    if risks:
        story.append(Paragraph("Risk Analysis", heading_style))
        
        # Risk summary table
        risk_data = [['Level', 'Count']]
        risk_counts = {"high": 0, "medium": 0, "low": 0}
        for risk in risks:
            level = risk.get("level", "Low").lower()
            if level in risk_counts:
                risk_counts[level] += 1
        
        for level, count in risk_counts.items():
            if count > 0:
                risk_data.append([level.upper(), str(count)])
        
        risk_table = Table(risk_data)
        risk_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 14),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(risk_table)
        story.append(Spacer(1, 20))
        
        # Individual risks
        for i, risk in enumerate(risks, 1):
            story.append(Paragraph(f"Risk #{i} - {risk.get('level', 'Low').upper()}", heading_style))
            story.append(Paragraph(f"<b>Description:</b> {risk.get('description', 'No description')}", normal_style))
            story.append(Paragraph(f"<b>Rationale:</b> {risk.get('rationale', 'No rationale')}", normal_style))
            story.append(Spacer(1, 12))
    else:
        story.append(Paragraph("Risk Analysis", heading_style))
        story.append(Paragraph("No significant risks identified.", normal_style))
        story.append(Spacer(1, 20))
    
    # Recommendations
    recommendations = report_data.get('recommendations', [])
    if recommendations:
        story.append(Paragraph("Recommendations", heading_style))
        for i, rec in enumerate(recommendations, 1):
            story.append(Paragraph(f"{i}. {rec}", normal_style))
        story.append(Spacer(1, 20))
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    
    # Return PDF response
    from fastapi.responses import Response
    return Response(
        content=buffer.getvalue(),
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=contract-review-{report_id}.pdf"}
    )

@app.get("/api/export/{report_id}/word")
def export_word(report_id: str):
    """Export report as Word document"""
    report_file = REPORTS_DIR / f"report-{report_id}.json"
    if not report_file.exists():
        raise HTTPException(status_code=404, detail="Report not found")
    
    with open(report_file, "r", encoding="utf-8") as f:
        report_data = json.load(f)
    
    # Create Word document
    doc = DocxDocument()
    
    # Title
    title = doc.add_heading('Contract Review Report', 0)
    title.alignment = 1  # Center alignment
    
    # Document info
    doc.add_paragraph(f"Document: {report_data.get('file_name', 'Unknown')}")
    doc.add_paragraph(f"Analysis Date: {report_data.get('timestamp', 'Unknown')}")
    doc.add_paragraph("")  # Empty line
    
    # Summary
    doc.add_heading('Executive Summary', level=1)
    summary_text = report_data.get('summary', 'No summary available')
    doc.add_paragraph(summary_text)
    doc.add_paragraph("")  # Empty line
    
    # Risk Analysis
    risks = report_data.get('risks', [])
    doc.add_heading('Risk Analysis', level=1)
    
    if risks:
        # Risk summary
        risk_counts = {"high": 0, "medium": 0, "low": 0}
        for risk in risks:
            level = risk.get("level", "Low").lower()
            if level in risk_counts:
                risk_counts[level] += 1
        
        doc.add_paragraph("Risk Summary:")
        for level, count in risk_counts.items():
            if count > 0:
                doc.add_paragraph(f"• {level.upper()}: {count} issue{'s' if count > 1 else ''}")
        
        doc.add_paragraph("")  # Empty line
        
        # Individual risks
        for i, risk in enumerate(risks, 1):
            doc.add_heading(f"Risk #{i} - {risk.get('level', 'Low').upper()}", level=2)
            doc.add_paragraph(f"Description: {risk.get('description', 'No description')}")
            doc.add_paragraph(f"Rationale: {risk.get('rationale', 'No rationale')}")
            doc.add_paragraph("")  # Empty line
    else:
        doc.add_paragraph("No significant risks identified.")
    
    # Recommendations
    recommendations = report_data.get('recommendations', [])
    if recommendations:
        doc.add_heading('Recommendations', level=1)
        for i, rec in enumerate(recommendations, 1):
            doc.add_paragraph(f"{i}. {rec}")
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Return Word response
    from fastapi.responses import Response
    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=contract-review-{report_id}.docx"}
    )

@app.get("/api/export/{report_id}/json")
def export_json(report_id: str):
    """Export report as JSON"""
    report_file = REPORTS_DIR / f"report-{report_id}.json"
    if not report_file.exists():
        raise HTTPException(status_code=404, detail="Report not found")
    
    with open(report_file, "r", encoding="utf-8") as f:
        report_data = json.load(f)
    
    # Return JSON response
    from fastapi.responses import Response
    return Response(
        content=json.dumps(report_data, ensure_ascii=False, indent=2),
        media_type="application/json",
        headers={"Content-Disposition": f"attachment; filename=contract-review-{report_id}.json"}
    )

# Remove:  @app.post("/api/review", response_model=ReviewResponse)
@app.post("/api/review")
async def api_review(file: UploadFile = File(...)):
    try:
        raw = await file.read()
        if not raw:
            raise HTTPException(status_code=400, detail="Empty file.")

        ext = Path(file.filename).suffix or ".pdf"
        upath = DATA_DIR / f"upload-{uuid.uuid4()}{ext}"
        upath.write_bytes(raw)
        input_hash = sha256_of_bytes(raw)

        text, kind = read_any_text(upath)
        print(f"DEBUG: Text extracted, length: {len(text)}")
        chunks = chunk_text(text)
        print(f"DEBUG: Text chunked into {len(chunks)} chunks")
        # Tag chunks with a per-upload doc_id so retrieval only hits this document
        doc_id = str(uuid.uuid4())
        upsert_chunks(chunks, doc_id)
        print("DEBUG: Chunks upserted to Qdrant")

        policy = Path(POLICY_FILE).read_text(encoding="utf-8")
        result = review_contract(text, policy, doc_id)
        result.update({"input_sha256": input_hash, "file_kind": kind})

        # Save report to reports directory
        report_id = str(uuid.uuid4())
        report_file = REPORTS_DIR / f"report-{report_id}.json"
        report_data = {
            **result,
            "report_id": report_id,
            "timestamp": datetime.now().isoformat(),
            "file_name": file.filename,
            "file_size": len(raw),
            "file_type": kind
        }
        report_file.write_text(json.dumps(report_data, ensure_ascii=False, indent=2), encoding="utf-8")
        
        # Add to history
        review_history.append({
            "id": report_id,
            "filename": file.filename,
            "timestamp": datetime.now().isoformat(),
            "summary": result.get("summary", ""),
            "risks_count": len(result.get("risks", [])),
            "file_size": len(raw)
        })
        
        return JSONResponse(status_code=200, content=report_data)

    except HTTPException as he:
        # clean JSON, no bytes
        return JSONResponse(status_code=he.status_code, content={"error": he.detail})
    except Exception as e:
        # last-resort error message (string only)
        print(f"ERROR in review endpoint: {str(e)}")
        print(f"ERROR type: {type(e)}")
        import traceback
        print(f"ERROR traceback: {traceback.format_exc()}")
        return JSONResponse(status_code=400, content={"error": str(e)})

def create_modern_ui():
    """Create a modern, production-ready UI using Gradio Blocks"""
    
    # Custom CSS for modern styling with enhanced features
    custom_css = """
    .gradio-container {
        max-width: 1400px !important;
        margin: 0 auto !important;
    }
    .upload-section {
        border: 2px dashed #e1e5e9;
        border-radius: 12px;
        padding: 2rem;
        text-align: center;
        background: #f8f9fa;
        transition: all 0.3s ease;
    }
    .upload-section:hover {
        border-color: #007bff;
        background: #e3f2fd;
    }
    .upload-section.collapsed {
        padding: 1rem;
        margin-bottom: 1rem;
    }
    .upload-section.collapsed .upload-content {
        display: none;
    }
    .upload-section.collapsed .collapse-toggle {
        display: block;
    }
    .collapse-toggle {
        display: none;
        background: #f8f9fa;
        border: 1px solid #dee2e6;
        border-radius: 8px;
        padding: 0.5rem 1rem;
        cursor: pointer;
        transition: all 0.3s ease;
    }
    .collapse-toggle:hover {
        background: #e9ecef;
    }
    .result-section {
        background: white;
        border-radius: 12px;
        padding: 1.5rem;
        box-shadow: 0 2px 8px rgba(0,0,0,0.1);
        margin-top: 1rem;
    }
    .risk-high { color: #dc3545; font-weight: bold; }
    .risk-medium { color: #fd7e14; font-weight: bold; }
    .risk-low { color: #28a745; font-weight: bold; }
    .risk-indicator {
        display: inline-block;
        width: 12px;
        height: 12px;
        border-radius: 50%;
        margin-right: 0.5rem;
    }
    .risk-indicator.high { background: #dc3545; }
    .risk-indicator.medium { background: #fd7e14; }
    .risk-indicator.low { background: #28a745; }
    .metric-card {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 1rem;
        border-radius: 8px;
        text-align: center;
    }
    .metric-value {
        font-size: 2rem;
        font-weight: bold;
        margin-bottom: 0.5rem;
    }
    .metric-label {
        font-size: 0.9rem;
        opacity: 0.9;
    }
    .progress-container {
        background: #f8f9fa;
        border-radius: 8px;
        padding: 1rem;
        margin: 1rem 0;
    }
    .progress-bar {
        width: 100%;
        height: 8px;
        background: #e9ecef;
        border-radius: 4px;
        overflow: hidden;
        margin: 0.5rem 0;
    }
    .progress-fill {
        height: 100%;
        background: linear-gradient(90deg, #007bff, #0056b3);
        border-radius: 4px;
        transition: width 0.3s ease;
        animation: pulse 2s infinite;
    }
    @keyframes pulse {
        0%, 100% { opacity: 1; }
        50% { opacity: 0.7; }
    }
    .status-message {
        text-align: center;
        color: #495057;
        font-style: italic;
        margin: 0.5rem 0;
    }
    .editable-text {
        background: transparent;
        border: 1px solid transparent;
        border-radius: 4px;
        padding: 0.5rem;
        transition: all 0.3s ease;
        min-height: 100px;
        resize: vertical;
    }
    .editable-text:hover {
        border-color: #dee2e6;
        background: #f8f9fa;
    }
    .editable-text:focus {
        border-color: #007bff;
        background: white;
        outline: none;
        box-shadow: 0 0 0 2px rgba(0, 123, 255, 0.25);
    }
    .comment-panel {
        background: #f8f9fa;
        border: 1px solid #dee2e6;
        border-radius: 8px;
        padding: 1rem;
        margin-top: 1rem;
    }
    .comment-item {
        background: white;
        border: 1px solid #e9ecef;
        border-radius: 6px;
        padding: 0.75rem;
        margin-bottom: 0.5rem;
    }
    .export-buttons {
        display: flex;
        gap: 0.5rem;
        margin-top: 1rem;
        flex-wrap: wrap;
    }
    .export-btn {
        background: #28a745;
        color: white;
        border: none;
        padding: 0.5rem 1rem;
        border-radius: 6px;
        cursor: pointer;
        font-size: 0.9rem;
        transition: all 0.3s ease;
    }
    .export-btn:hover {
        background: #218838;
        transform: translateY(-1px);
    }
    .export-btn.secondary {
        background: #6c757d;
    }
    .export-btn.secondary:hover {
        background: #5a6268;
    }
    """
    
    def process_document(file, progress=gr.Progress()):
        """Process uploaded document with progress tracking"""
        try:
            print(f"DEBUG: process_document invoked. file={file} name={getattr(file, 'name', None)}")
        except Exception:
            pass
        if file is None:
            return None, "Please upload a document first.", "", "", "", "", False, ""
        
        try:
            progress(0.1, desc="📤 Uploading document...")
            
            mime = "application/pdf" if file.name.lower().endswith(".pdf") else \
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            
            progress(0.2, desc="🔍 Extracting text from document...")
            
            with open(file.name, "rb") as fh:
                progress(0.3, desc="🧠 Analyzing document with AI...")
                response = requests.post("http://localhost:7860/api/review",
                                       files={"file": (Path(file.name).name, fh, mime)},
                            timeout=300)
            
            progress(0.8, desc="📊 Processing results...")
            
            if response.status_code == 200:
                result = response.json()
                progress(1.0, desc="✅ Analysis complete!")
                return (format_results(result), 
                       create_summary_html(result), 
                       create_key_terms_html(result),
                       create_risks_html(result), 
                       create_recommendations_html(result), 
                       result.get("report_id", ""),
                       True,  # Show results
                       file.name)  # Return filename for display
            else:
                progress(1.0, desc="❌ Analysis failed")
                return None, f"Error: {response.status_code} - {response.text[:500]}", "", "", "", "", False, ""
                
        except Exception as e:
            progress(1.0, desc="❌ Analysis failed")
            return None, f"Error processing document: {str(e)}", "", "", "", "", False, ""
    
    def format_results(result):
        """Format results as JSON"""
        return json.dumps(result, ensure_ascii=False, indent=2)
    
    def create_key_points_html(key_points):
        """Create HTML for key points section"""
        if not key_points:
            return ""
        
        points_html = ""
        for i, point in enumerate(key_points, 1):
            points_html += f'<div class="key-point-item"><span class="point-number">{i}.</span> {point}</div>'
        
        return f"""
        <div class="summary-section">
            <h3>🔑 Key Points</h3>
            <div class="summary-content">
                <div class="key-points-list">
                    {points_html}
                </div>
            </div>
        </div>
        """

    def create_key_terms_html(data):
        """Create HTML for key terms section"""
        key_points = data.get("key_points", [])
        citations = data.get("citations", [])
        
        if not key_points and not citations:
            return """
            <div style="text-align: center; color: #6c757d; padding: 2rem;">
                <h3>🔑 Key Terms</h3>
                <p>No key terms identified in this document.</p>
            </div>
            """
        
        html = """
        <div style="max-width: 100%; margin: 0 auto;">
            <h3 style="color: #2c3e50; margin-bottom: 1.5rem; text-align: center;">🔑 Key Terms & Clauses</h3>
        """
        
        if key_points:
            html += """
            <div style="margin-bottom: 2rem;">
                <h4 style="color: #34495e; margin-bottom: 1rem; border-bottom: 2px solid #3498db; padding-bottom: 0.5rem;">
                    📋 Document Sections
                </h4>
                <div style="background: #f8f9fa; border-radius: 8px; padding: 1rem;">
            """
            for i, point in enumerate(key_points, 1):
                html += f"""
                <div style="background: white; margin-bottom: 0.75rem; padding: 1rem; border-radius: 6px; border-left: 4px solid #3498db; box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
                    <div style="display: flex; align-items: flex-start;">
                        <span style="background: #3498db; color: white; border-radius: 50%; width: 24px; height: 24px; display: flex; align-items: center; justify-content: center; font-size: 0.8rem; font-weight: bold; margin-right: 0.75rem; flex-shrink: 0;">{i}</span>
                        <div style="flex: 1; line-height: 1.5; color: #2c3e50;">{point}</div>
                    </div>
                </div>
                """
            html += "</div></div>"
        
        if citations:
            html += """
            <div>
                <h4 style="color: #34495e; margin-bottom: 1rem; border-bottom: 2px solid #e74c3c; padding-bottom: 0.5rem;">
                    📖 Document Citations
                </h4>
                <div style="background: #fdf2f2; border-radius: 8px; padding: 1rem;">
            """
            for i, citation in enumerate(citations, 1):
                section = citation.get("section", "Unknown Section")
                text = citation.get("text", "No text available")
                html += f"""
                <div style="background: white; margin-bottom: 0.75rem; padding: 1rem; border-radius: 6px; border-left: 4px solid #e74c3c; box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
                    <div style="font-weight: 600; color: #c0392b; margin-bottom: 0.5rem;">{section}</div>
                    <div style="color: #2c3e50; line-height: 1.5; font-style: italic;">"{text}"</div>
                </div>
                """
            html += "</div></div>"
        
        html += "</div>"
        return html

    def create_summary_html(result, editable_summary=None):
        """Create HTML summary with editable functionality"""
        summary = editable_summary if editable_summary else result.get("summary", "No summary available")
        doc_type = result.get("document_type", "Unknown")
        key_points = result.get("key_points", [])
        report_id = result.get("report_id", "")
        
        # Clean and format the summary text
        def format_text(text, max_length=500):
            """Format text for better readability"""
            if len(text) <= max_length:
                return text
            
            # Find the last complete sentence within the limit
            truncated = text[:max_length]
            last_period = truncated.rfind('.')
            last_newline = truncated.rfind('\n')
            
            # Use the last complete sentence or paragraph
            if last_period > max_length * 0.7:
                return text[:last_period + 1] + "..."
            elif last_newline > max_length * 0.7:
                return text[:last_newline] + "..."
            else:
                return text[:max_length] + "..."
        
        # Format the summary with proper line breaks and structure
        formatted_summary = format_text(summary)
        
        # Add line breaks for better readability
        formatted_summary = formatted_summary.replace('. ', '.<br>')
        formatted_summary = formatted_summary.replace('\n', '<br>')
        
        # Determine document type with better formatting
        type_display = {
            "Contract": "📋 Contract",
            "Technical Document": "📊 Technical Document", 
            "Policy": "📜 Policy",
            "Other": "📄 Document"
        }.get(doc_type, f"📄 {doc_type}")
        
        html = f"""
        <div class="result-section" data-report-id="{report_id}">
            <div style="display: flex; align-items: center; justify-content: space-between; margin-bottom: 1.5rem;">
                <h3 style="margin: 0; color: #2c3e50;">📄 Document Summary</h3>
                <div style="display: flex; gap: 0.5rem;">
                    <button onclick="toggleEditMode()" style="background: #007bff; color: white; border: none; padding: 0.5rem 1rem; border-radius: 6px; cursor: pointer; font-size: 0.9rem;">
                        ✏️ Edit Summary
                    </button>
                </div>
            </div>
            
            <div style="background: #f8f9fa; padding: 1rem; border-radius: 8px; margin-bottom: 1.5rem;">
                <div style="display: flex; align-items: center; margin-bottom: 0.5rem;">
                    <span style="font-weight: 600; color: #495057; margin-right: 0.5rem;">Document Type:</span>
                    <span style="background: #e3f2fd; color: #1565c0; padding: 0.25rem 0.75rem; border-radius: 20px; font-size: 0.9rem; font-weight: 500;">
                        {type_display}
                    </span>
                </div>
            </div>
            
            <div style="margin-bottom: 1.5rem;">
                <h4 style="color: #34495e; margin-bottom: 0.75rem; font-size: 1.1rem;">📝 Executive Summary</h4>
                <div id="summary-display" style="background: white; padding: 1.25rem; border-left: 4px solid #3498db; border-radius: 0 8px 8px 0; line-height: 1.6; color: #2c3e50;">
                    {formatted_summary}
                </div>
                <textarea id="summary-edit" class="editable-text" style="display: none; width: 100%; min-height: 120px; padding: 1.25rem; border: 1px solid #3498db; border-radius: 0 8px 8px 0; font-family: inherit; font-size: inherit; line-height: 1.6; color: #2c3e50; background: white; resize: vertical;">{summary}</textarea>
                <div id="edit-controls" style="display: none; margin-top: 1rem; text-align: right;">
                    <button onclick="saveSummary()" style="background: #28a745; color: white; border: none; padding: 0.5rem 1rem; border-radius: 6px; cursor: pointer; font-size: 0.9rem; margin-right: 0.5rem;">
                        💾 Save Changes
                    </button>
                    <button onclick="cancelEdit()" style="background: #6c757d; color: white; border: none; padding: 0.5rem 1rem; border-radius: 6px; cursor: pointer; font-size: 0.9rem;">
                        ❌ Cancel
                    </button>
                </div>
            </div>
        """
        
        if key_points:
            html += f"""
            <div>
                <h4 style="color: #34495e; margin-bottom: 0.75rem; font-size: 1.1rem;">🎯 Key Points</h4>
                <div style="background: #fff3cd; border: 1px solid #ffeaa7; border-radius: 8px; padding: 1rem;">
                    <ul style="margin: 0; padding-left: 1.5rem; color: #856404;">
            """
            for point in key_points:
                # Format key points for better readability
                formatted_point = point.replace('. ', '.<br>') if len(point) > 100 else point
                html += f"<li style='margin-bottom: 0.5rem; line-height: 1.5;'>{formatted_point}</li>"
            html += """
                    </ul>
                </div>
            </div>
            """
        
        # Add export buttons
        html += f"""
            <div class="export-buttons">
                <button onclick="exportToPDF()" class="export-btn">
                    📄 Export PDF
                </button>
                <button onclick="exportToWord()" class="export-btn">
                    📝 Export Word
                </button>
                <button onclick="exportToJSON()" class="export-btn secondary">
                    💾 Export JSON
                </button>
            </div>
        """
        
        # Add JavaScript for edit functionality
        html += """
            <script>
                function toggleEditMode() {
                    const display = document.getElementById('summary-display');
                    const edit = document.getElementById('summary-edit');
                    const controls = document.getElementById('edit-controls');
                    
                    if (display.style.display !== 'none') {
                        display.style.display = 'none';
                        edit.style.display = 'block';
                        controls.style.display = 'block';
                        edit.focus();
                    } else {
                        display.style.display = 'block';
                        edit.style.display = 'none';
                        controls.style.display = 'none';
                    }
                }
                
                function saveSummary() {
                    const edit = document.getElementById('summary-edit');
                    const display = document.getElementById('summary-display');
                    const newText = edit.value;
                    
                    // Format the text for display
                    const formatted = newText.replace(/\. /g, '.<br>').replace(/\\n/g, '<br>');
                    display.innerHTML = formatted;
                    
                    // Hide edit mode
                    toggleEditMode();
                    
                    // Show success message
                    showNotification('Summary updated successfully!', 'success');
                }
                
                function cancelEdit() {
                    const display = document.getElementById('summary-display');
                    const edit = document.getElementById('summary-edit');
                    const controls = document.getElementById('edit-controls');
                    
                    display.style.display = 'block';
                    edit.style.display = 'none';
                    controls.style.display = 'none';
                }
                
                function exportToPDF() {
                    const reportId = document.querySelector('[data-report-id]')?.getAttribute('data-report-id');
                    if (reportId) {
                        window.open(`/api/export/${reportId}/pdf`, '_blank');
                        showNotification('PDF export started!', 'success');
                    } else {
                        showNotification('No report available for export', 'error');
                    }
                }
                
                function exportToWord() {
                    const reportId = document.querySelector('[data-report-id]')?.getAttribute('data-report-id');
                    if (reportId) {
                        window.open(`/api/export/${reportId}/word`, '_blank');
                        showNotification('Word export started!', 'success');
                    } else {
                        showNotification('No report available for export', 'error');
                    }
                }
                
                function exportToJSON() {
                    const reportId = document.querySelector('[data-report-id]')?.getAttribute('data-report-id');
                    if (reportId) {
                        window.open(`/api/export/${reportId}/json`, '_blank');
                        showNotification('JSON export started!', 'success');
                    } else {
                        showNotification('No report available for export', 'error');
                    }
                }
                
                function showNotification(message, type) {
                    const notification = document.createElement('div');
                    notification.style.cssText = `
                        position: fixed;
                        top: 20px;
                        right: 20px;
                        padding: 1rem 1.5rem;
                        border-radius: 8px;
                        color: white;
                        font-weight: 500;
                        z-index: 1000;
                        animation: slideIn 0.3s ease;
                    `;
                    
                    const colors = {
                        success: '#28a745',
                        info: '#17a2b8',
                        warning: '#ffc107',
                        error: '#dc3545'
                    };
                    
                    notification.style.background = colors[type] || colors.info;
                    notification.textContent = message;
                    
                    document.body.appendChild(notification);
                    
                    setTimeout(() => {
                        notification.style.animation = 'slideOut 0.3s ease';
                        setTimeout(() => notification.remove(), 300);
                    }, 3000);
                }
                
                // Add CSS animations
                const style = document.createElement('style');
                style.textContent = `
                    @keyframes slideIn {
                        from { transform: translateX(100%); opacity: 0; }
                        to { transform: translateX(0); opacity: 1; }
                    }
                    @keyframes slideOut {
                        from { transform: translateX(0); opacity: 1; }
                        to { transform: translateX(100%); opacity: 0; }
                    }
                `;
                document.head.appendChild(style);
            </script>
        """
        
        html += "</div>"
        return html
    
    def create_risks_html(result):
        """Create HTML for risks with enhanced traffic light visualization"""
        risks = result.get("risks", [])
        if not risks:
            return """
            <div class="result-section">
                <div style="display: flex; align-items: center; margin-bottom: 1.5rem;">
                    <h3 style="margin: 0; color: #2c3e50;">⚠️ Risk Analysis</h3>
                    <div style="margin-left: auto; display: flex; align-items: center;">
                        <div class="risk-indicator low"></div>
                        <span style="color: #28a745; font-weight: 600; font-size: 0.9rem;">LOW RISK</span>
                    </div>
                </div>
                <div style="background: #d4edda; border: 1px solid #c3e6cb; border-radius: 8px; padding: 1.5rem; text-align: center;">
                    <div style="color: #155724; font-size: 1.1rem; font-weight: 500; margin-bottom: 0.5rem;">
                        🟢 No significant risks identified
                    </div>
                    <div style="color: #6c757d; font-size: 0.9rem;">
                        This document appears to be low risk based on the analysis.
                    </div>
                </div>
            </div>
            """
        
        # Count risks by level and determine overall risk level
        risk_counts = {"high": 0, "medium": 0, "low": 0}
        for risk in risks:
            level = risk.get("level", "Low").lower()
            if level in risk_counts:
                risk_counts[level] += 1
        
        # Determine overall risk level
        if risk_counts["high"] > 0:
            overall_risk = "HIGH"
            overall_color = "#dc3545"
            overall_indicator = "high"
            overall_emoji = "🔴"
        elif risk_counts["medium"] > 0:
            overall_risk = "MEDIUM"
            overall_color = "#fd7e14"
            overall_indicator = "medium"
            overall_emoji = "🟡"
        else:
            overall_risk = "LOW"
            overall_color = "#28a745"
            overall_indicator = "low"
            overall_emoji = "🟢"
        
        html = f"""
        <div class="result-section">
            <div style="display: flex; align-items: center; justify-content: space-between; margin-bottom: 1.5rem;">
                <h3 style="margin: 0; color: #2c3e50;">⚠️ Risk Analysis</h3>
                <div style="display: flex; align-items: center; gap: 1rem;">
                    <div style="display: flex; align-items: center;">
                        <div class="risk-indicator {overall_indicator}"></div>
                        <span style="color: {overall_color}; font-weight: 600; font-size: 0.9rem;">{overall_emoji} {overall_risk} RISK</span>
                    </div>
                    <div style="display: flex; gap: 0.5rem;">
        """
        
        # Add risk level badges with traffic light colors
        for level, count in risk_counts.items():
            if count > 0:
                colors = {
                    "high": {"bg": "#f8d7da", "text": "#721c24", "border": "#dc3545", "emoji": "🔴"},
                    "medium": {"bg": "#fff3cd", "text": "#856404", "border": "#fd7e14", "emoji": "🟡"},
                    "low": {"bg": "#d4edda", "text": "#155724", "border": "#28a745", "emoji": "🟢"}
                }
                color = colors[level]
                html += f"""
                <span style="background: {color['bg']}; color: {color['text']}; border: 1px solid {color['border']}; 
                           padding: 0.25rem 0.75rem; border-radius: 20px; font-size: 0.8rem; font-weight: 500;">
                    {color['emoji']} {level.upper()}: {count}
                </span>
                """
        
        html += """
                    </div>
                </div>
            </div>
        """
        
        # Add risk summary bar
        total_risks = sum(risk_counts.values())
        html += f"""
            <div style="background: #f8f9fa; border: 1px solid #dee2e6; border-radius: 8px; padding: 1rem; margin-bottom: 1.5rem;">
                <div style="display: flex; justify-content: space-between; align-items: center; margin-bottom: 0.5rem;">
                    <span style="font-weight: 600; color: #495057;">Risk Summary</span>
                    <span style="color: #6c757d; font-size: 0.9rem;">{total_risks} issue{'' if total_risks == 1 else 's'} found</span>
                </div>
                <div style="display: flex; gap: 0.5rem; align-items: center;">
                    <div style="flex: 1; height: 8px; background: #e9ecef; border-radius: 4px; overflow: hidden;">
                        <div style="height: 100%; background: #dc3545; width: {risk_counts['high']/total_risks*100 if total_risks > 0 else 0}%; border-radius: 4px;"></div>
                    </div>
                    <div style="flex: 1; height: 8px; background: #e9ecef; border-radius: 4px; overflow: hidden;">
                        <div style="height: 100%; background: #fd7e14; width: {risk_counts['medium']/total_risks*100 if total_risks > 0 else 0}%; border-radius: 4px;"></div>
                    </div>
                    <div style="flex: 1; height: 8px; background: #e9ecef; border-radius: 4px; overflow: hidden;">
                        <div style="height: 100%; background: #28a745; width: {risk_counts['low']/total_risks*100 if total_risks > 0 else 0}%; border-radius: 4px;"></div>
                    </div>
                </div>
                <div style="display: flex; justify-content: space-between; margin-top: 0.5rem; font-size: 0.8rem; color: #6c757d;">
                    <span>🔴 High</span>
                    <span>🟡 Medium</span>
                    <span>🟢 Low</span>
                </div>
            </div>
        """
        
        for i, risk in enumerate(risks, 1):
            level = risk.get("level", "Low").lower()
            description = risk.get("description", "No description provided")
            rationale = risk.get("rationale", "No rationale provided")
            
            # Format text for better readability
            description = description.replace('. ', '.<br>') if len(description) > 150 else description
            rationale = rationale.replace('. ', '.<br>') if len(rationale) > 150 else rationale
            
            colors = {
                "high": {"border": "#dc3545", "bg": "#f8d7da", "badge": "#dc3545", "emoji": "🔴"},
                "medium": {"border": "#fd7e14", "bg": "#fff3cd", "badge": "#fd7e14", "emoji": "🟡"},
                "low": {"border": "#28a745", "bg": "#d4edda", "badge": "#28a745", "emoji": "🟢"}
            }
            color = colors.get(level, colors["low"])
            
            html += f"""
            <div style="margin-bottom: 1.25rem; border: 1px solid {color['border']}; border-radius: 8px; 
                        background: {color['bg']}; overflow: hidden;">
                <div style="background: {color['badge']}; color: white; padding: 0.75rem 1rem; font-weight: 600; 
                            font-size: 0.9rem; text-transform: uppercase; letter-spacing: 0.5px; display: flex; align-items: center;">
                    <span style="margin-right: 0.5rem;">{color['emoji']}</span>
                    Risk #{i} • {level.upper()} Priority
                </div>
                <div style="padding: 1rem;">
                    <div style="margin-bottom: 0.75rem;">
                        <div style="font-weight: 600; color: #2c3e50; margin-bottom: 0.25rem; font-size: 1rem;">
                            📋 Issue Description
                        </div>
                        <div style="color: #495057; line-height: 1.5;">
                            {description}
                        </div>
                    </div>
                    <div>
                        <div style="font-weight: 600; color: #2c3e50; margin-bottom: 0.25rem; font-size: 1rem;">
                            💡 Rationale
                        </div>
                        <div style="color: #6c757d; line-height: 1.5; font-style: italic;">
                            {rationale}
                        </div>
                    </div>
                </div>
            </div>
            """
        
        html += "</div>"
        return html
    
    def create_recommendations_html(result):
        """Create HTML for recommendations with smart suggestions"""
        recommendations = result.get("recommendations", [])
        citations = result.get("citations", [])
        risks = result.get("risks", [])
        summary = result.get("summary", "")
        
        # Generate smart recommendations based on analysis
        smart_recommendations = generate_smart_recommendations(risks, summary, recommendations)
        
        html = f"""
        <div class="result-section">
            <div style="display: flex; align-items: center; margin-bottom: 1.5rem;">
                <h3 style="margin: 0; color: #2c3e50;">💡 Smart Recommendations & Citations</h3>
            </div>
        """
        
        # Display smart recommendations
        if smart_recommendations:
            html += f"""
            <div style="margin-bottom: 2rem;">
                <h4 style="color: #34495e; margin-bottom: 1rem; font-size: 1.1rem; display: flex; align-items: center;">
                    🤖 AI-Generated Recommendations ({len(smart_recommendations)})
                </h4>
                <div style="background: linear-gradient(135deg, #e3f2fd 0%, #bbdefb 100%); border: 1px solid #2196f3; border-radius: 8px; padding: 1rem;">
            """
            
            for i, rec in enumerate(smart_recommendations, 1):
                priority = rec.get("priority", "medium")
                category = rec.get("category", "General")
                
                # Priority colors
                priority_colors = {
                    "high": {"bg": "#ffebee", "border": "#f44336", "badge": "#f44336"},
                    "medium": {"bg": "#fff3e0", "border": "#ff9800", "badge": "#ff9800"},
                    "low": {"bg": "#e8f5e8", "border": "#4caf50", "badge": "#4caf50"}
                }
                colors = priority_colors.get(priority, priority_colors["medium"])
                
                html += f"""
                <div style="margin-bottom: 1rem; padding: 0.75rem; background: white; border-radius: 6px; 
                            border-left: 4px solid {colors['border']}; box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
                    <div style="display: flex; align-items: flex-start;">
                        <span style="background: {colors['badge']}; color: white; width: 24px; height: 24px; border-radius: 50%; 
                                     display: flex; align-items: center; justify-content: center; font-size: 0.8rem; 
                                     font-weight: 600; margin-right: 0.75rem; flex-shrink: 0; margin-top: 0.125rem;">
                            {i}
                        </span>
                        <div style="flex: 1;">
                            <div style="display: flex; justify-content: space-between; align-items: center; margin-bottom: 0.5rem;">
                                <span style="background: {colors['badge']}; color: white; padding: 0.25rem 0.5rem; border-radius: 12px; font-size: 0.8rem; font-weight: 500;">
                                    {category}
                                </span>
                                <span style="color: {colors['border']}; font-size: 0.8rem; font-weight: 600; text-transform: uppercase;">
                                    {priority} Priority
                                </span>
                            </div>
                            <div style="color: #2c3e50; line-height: 1.5; margin-bottom: 0.5rem;">
                                {rec.get('recommendation', '')}
                            </div>
                            <div style="color: #6c757d; font-size: 0.9rem; font-style: italic;">
                                💡 {rec.get('rationale', '')}
                            </div>
                        </div>
                    </div>
                </div>
                """
            
            html += """
                </div>
            </div>
            """
        
        # Display original recommendations if any
        if recommendations:
            html += f"""
            <div style="margin-bottom: 2rem;">
                <h4 style="color: #34495e; margin-bottom: 1rem; font-size: 1.1rem; display: flex; align-items: center;">
                    📋 Original Recommendations ({len(recommendations)})
                </h4>
                <div style="background: #e8f5e8; border: 1px solid #c3e6cb; border-radius: 8px; padding: 1rem;">
            """
            
            for i, rec in enumerate(recommendations, 1):
                # Format recommendations for better readability
                formatted_rec = rec.replace('. ', '.<br>') if len(rec) > 100 else rec
                html += f"""
                <div style="margin-bottom: 1rem; padding: 0.75rem; background: white; border-radius: 6px; 
                            border-left: 4px solid #28a745;">
                    <div style="display: flex; align-items: flex-start;">
                        <span style="background: #28a745; color: white; width: 24px; height: 24px; border-radius: 50%; 
                                     display: flex; align-items: center; justify-content: center; font-size: 0.8rem; 
                                     font-weight: 600; margin-right: 0.75rem; flex-shrink: 0; margin-top: 0.125rem;">
                            {i}
                        </span>
                        <div style="color: #2c3e50; line-height: 1.5;">
                            {formatted_rec}
                        </div>
                    </div>
                </div>
                """
            
            html += """
                </div>
            </div>
            """
        
        # Show message if no recommendations
        if not smart_recommendations and not recommendations:
            html += """
            <div style="background: #f8f9fa; border: 1px solid #dee2e6; border-radius: 8px; padding: 1.5rem; 
                        text-align: center; margin-bottom: 2rem;">
                <div style="color: #6c757d; font-size: 1rem;">
                    ℹ️ No specific recommendations provided for this document.
                </div>
            </div>
            """
        
        if citations:
            html += f"""
            <div>
                <h4 style="color: #34495e; margin-bottom: 1rem; font-size: 1.1rem; display: flex; align-items: center;">
                    📚 Document Citations ({len(citations)})
                </h4>
                <div style="background: #fff3cd; border: 1px solid #ffeaa7; border-radius: 8px; padding: 1rem;">
            """
            
            for i, citation in enumerate(citations, 1):
                section = citation.get('section', 'Unknown Section')
                text = citation.get('text', 'No text available')
                
                # Format citation text for better readability
                formatted_text = text.replace('. ', '.<br>') if len(text) > 120 else text
                
                html += f"""
                <div style="margin-bottom: 1rem; padding: 0.75rem; background: white; border-radius: 6px; 
                            border-left: 4px solid #ffc107;">
                    <div style="display: flex; align-items: flex-start;">
                        <span style="background: #ffc107; color: #212529; width: 24px; height: 24px; border-radius: 50%; 
                                     display: flex; align-items: center; justify-content: center; font-size: 0.8rem; 
                                     font-weight: 600; margin-right: 0.75rem; flex-shrink: 0; margin-top: 0.125rem;">
                            {i}
                        </span>
                        <div>
                            <div style="font-weight: 600; color: #856404; margin-bottom: 0.25rem; font-size: 0.9rem;">
                                📍 {section}
                            </div>
                            <div style="color: #495057; line-height: 1.5; font-style: italic;">
                                "{formatted_text}"
                            </div>
                        </div>
                    </div>
                </div>
                """
            
            html += """
                </div>
            </div>
            """
        else:
            html += """
            <div style="background: #f8f9fa; border: 1px solid #dee2e6; border-radius: 8px; padding: 1.5rem; 
                        text-align: center;">
                <div style="color: #6c757d; font-size: 1rem;">
                    ℹ️ No specific citations found in this document.
                </div>
            </div>
            """
        
        html += "</div>"
        return html
    
    def create_risk_overview(result):
        """Create a compact risk overview for the top of results"""
        risks = result.get("risks", [])
        if not risks:
            return """
            <div style="background: linear-gradient(135deg, #d4edda 0%, #c3e6cb 100%); border: 1px solid #28a745; border-radius: 12px; padding: 1.5rem; text-align: center; margin-bottom: 1.5rem;">
                <div style="display: flex; align-items: center; justify-content: center; margin-bottom: 0.5rem;">
                    <div class="risk-indicator low" style="margin-right: 0.5rem;"></div>
                    <h3 style="margin: 0; color: #155724; font-size: 1.2rem;">🟢 LOW RISK DOCUMENT</h3>
                </div>
                <p style="margin: 0; color: #155724; font-size: 1rem;">No significant risks identified in this document.</p>
            </div>
            """
        
        # Count risks by level
        risk_counts = {"high": 0, "medium": 0, "low": 0}
        for risk in risks:
            level = risk.get("level", "Low").lower()
            if level in risk_counts:
                risk_counts[level] += 1
        
        # Determine overall risk level
        if risk_counts["high"] > 0:
            overall_risk = "HIGH RISK"
            overall_color = "#dc3545"
            overall_bg = "linear-gradient(135deg, #f8d7da 0%, #f5c6cb 100%)"
            overall_border = "#dc3545"
            overall_emoji = "🔴"
        elif risk_counts["medium"] > 0:
            overall_risk = "MEDIUM RISK"
            overall_color = "#fd7e14"
            overall_bg = "linear-gradient(135deg, #fff3cd 0%, #ffeaa7 100%)"
            overall_border = "#fd7e14"
            overall_emoji = "🟡"
        else:
            overall_risk = "LOW RISK"
            overall_color = "#28a745"
            overall_bg = "linear-gradient(135deg, #d4edda 0%, #c3e6cb 100%)"
            overall_border = "#28a745"
            overall_emoji = "🟢"
        
        total_risks = sum(risk_counts.values())
        
        return f"""
        <div style="background: {overall_bg}; border: 2px solid {overall_border}; border-radius: 12px; padding: 1.5rem; margin-bottom: 1.5rem;">
            <div style="display: flex; align-items: center; justify-content: space-between; margin-bottom: 1rem;">
                <div style="display: flex; align-items: center;">
                    <div class="risk-indicator {overall_risk.lower().split()[0]}" style="margin-right: 0.5rem;"></div>
                    <h3 style="margin: 0; color: {overall_color}; font-size: 1.2rem;">{overall_emoji} {overall_risk}</h3>
                </div>
                <div style="background: rgba(255,255,255,0.8); padding: 0.5rem 1rem; border-radius: 20px; font-weight: 600; color: {overall_color};">
                    {total_risks} Issue{'' if total_risks == 1 else 's'} Found
                </div>
            </div>
            
            <div style="display: flex; gap: 1rem; justify-content: center;">
                <div style="text-align: center;">
                    <div style="font-size: 1.5rem; font-weight: bold; color: #dc3545;">{risk_counts['high']}</div>
                    <div style="font-size: 0.9rem; color: #721c24;">🔴 High</div>
                </div>
                <div style="text-align: center;">
                    <div style="font-size: 1.5rem; font-weight: bold; color: #fd7e14;">{risk_counts['medium']}</div>
                    <div style="font-size: 0.9rem; color: #856404;">🟡 Medium</div>
                </div>
                <div style="text-align: center;">
                    <div style="font-size: 1.5rem; font-weight: bold; color: #28a745;">{risk_counts['low']}</div>
                    <div style="font-size: 0.9rem; color: #155724;">🟢 Low</div>
                </div>
            </div>
        </div>
        """
    
    # Global comments storage (in production, this would be a database)
    comments_storage = {}
    
    def create_comments_html(comments_list):
        """Create HTML for displaying comments"""
        if not comments_list:
            return "<div style='text-align: center; color: #6c757d; padding: 2rem;'>No comments yet. Add your first comment below!</div>"
        
        html = "<div class='comment-panel'>"
        for i, comment in enumerate(comments_list):
            timestamp = comment.get('timestamp', 'Unknown time')
            section = comment.get('section', 'General')
            text = comment.get('text', '')
            safe_text = text.replace("\n", "<br>")
            author = comment.get('author', 'Anonymous')
            
            # Format timestamp
            try:
                from datetime import datetime
                dt = datetime.fromisoformat(timestamp.replace('Z', '+00:00'))
                formatted_time = dt.strftime('%Y-%m-%d %H:%M')
            except:
                formatted_time = timestamp
            
            # Section color coding
            section_colors = {
                "Summary": "#3498db",
                "Risk Analysis": "#e74c3c", 
                "Recommendations": "#f39c12",
                "General": "#95a5a6"
            }
            section_color = section_colors.get(section, "#95a5a6")
            
            html += f"""
            <div class="comment-item">
                <div style="display: flex; justify-content: space-between; align-items: center; margin-bottom: 0.5rem;">
                    <div style="display: flex; align-items: center; gap: 0.5rem;">
                        <span style="background: {section_color}; color: white; padding: 0.25rem 0.5rem; border-radius: 12px; font-size: 0.8rem; font-weight: 500;">
                            {section}
                        </span>
                        <span style="font-weight: 600; color: #2c3e50;">{author}</span>
                    </div>
                    <span style="color: #6c757d; font-size: 0.8rem;">{formatted_time}</span>
                </div>
                <div style="color: #495057; line-height: 1.5; margin-bottom: 0.5rem;">
                    {safe_text}
                </div>
                <div style="text-align: right;">
                    <button onclick="deleteComment({i})" style="background: #dc3545; color: white; border: none; padding: 0.25rem 0.5rem; border-radius: 4px; cursor: pointer; font-size: 0.8rem;">
                        🗑️ Delete
                    </button>
                </div>
            </div>
            """
        
        html += "</div>"
        
        # Add JavaScript for comment management
        html += """
        <script>
            function deleteComment(index) {
                if (confirm('Are you sure you want to delete this comment?')) {
                    // This would trigger a server-side delete in a real implementation
                    showNotification('Comment deleted successfully!', 'success');
                    // Refresh comments display
                    setTimeout(() => location.reload(), 1000);
                }
            }
        </script>
        """
        
        return html
    
    def add_comment(comment_text, section, report_id):
        """Add a new comment"""
        if not comment_text.strip():
            return "Comment cannot be empty.", gr.update()
        
        if report_id not in comments_storage:
            comments_storage[report_id] = []
        
        from datetime import datetime
        new_comment = {
            'text': comment_text,
            'section': section,
            'timestamp': datetime.now().isoformat(),
            'author': 'User'  # In production, this would be the logged-in user
        }
        
        comments_storage[report_id].append(new_comment)
        
        # Return updated comments HTML
        updated_comments_html = create_comments_html(comments_storage[report_id])
        
        return "Comment added successfully!", updated_comments_html
    
    def generate_smart_recommendations(risks, summary, existing_recommendations):
        """Generate smart recommendations based on analysis results"""
        recommendations = []
        
        # Analyze risks and generate recommendations
        risk_analysis = {
            "liability": False,
            "indemnity": False,
            "confidentiality": False,
            "termination": False,
            "governing_law": False,
            "intellectual_property": False,
            "data_protection": False,
            "assignment": False
        }
        
        # Check for common risk patterns
        summary_lower = summary.lower()
        for risk in risks:
            risk_desc = risk.get("description", "").lower()
            risk_rationale = risk.get("rationale", "").lower()
            combined_text = f"{risk_desc} {risk_rationale}"
            
            if any(term in combined_text for term in ["liability", "damages", "loss"]):
                risk_analysis["liability"] = True
            if any(term in combined_text for term in ["indemnify", "indemnity", "hold harmless"]):
                risk_analysis["indemnity"] = True
            if any(term in combined_text for term in ["confidential", "proprietary", "non-disclosure"]):
                risk_analysis["confidentiality"] = True
            if any(term in combined_text for term in ["terminate", "termination", "end"]):
                risk_analysis["termination"] = True
            if any(term in combined_text for term in ["governing law", "jurisdiction", "venue"]):
                risk_analysis["governing_law"] = True
            if any(term in combined_text for term in ["intellectual property", "copyright", "patent", "trademark"]):
                risk_analysis["intellectual_property"] = True
            if any(term in combined_text for term in ["data", "privacy", "gdpr", "personal information"]):
                risk_analysis["data_protection"] = True
            if any(term in combined_text for term in ["assign", "assignment", "transfer"]):
                risk_analysis["assignment"] = True
        
        # Generate recommendations based on analysis
        if risk_analysis["liability"]:
            recommendations.append({
                "recommendation": "Consider adding liability limitation clauses to cap potential damages",
                "rationale": "Liability risks were identified in the analysis. Limiting liability can protect your organization from excessive financial exposure.",
                "category": "Risk Mitigation",
                "priority": "high"
            })
        
        if risk_analysis["indemnity"]:
            recommendations.append({
                "recommendation": "Review indemnity clauses to ensure mutual protection and fair allocation of risks",
                "rationale": "Indemnity provisions were found. Ensure the scope and limits are reasonable and reciprocal where appropriate.",
                "category": "Contract Terms",
                "priority": "high"
            })
        
        if not risk_analysis["confidentiality"]:
            recommendations.append({
                "recommendation": "Consider adding confidentiality provisions to protect sensitive information",
                "rationale": "No confidentiality clauses were identified. This may be necessary depending on the nature of the agreement.",
                "category": "Missing Clauses",
                "priority": "medium"
            })
        
        if not risk_analysis["termination"]:
            recommendations.append({
                "recommendation": "Add clear termination provisions including notice periods and consequences",
                "rationale": "Termination clauses provide clarity on how the agreement can be ended, protecting both parties.",
                "category": "Missing Clauses",
                "priority": "medium"
            })
        
        if not risk_analysis["governing_law"]:
            recommendations.append({
                "recommendation": "Specify governing law and jurisdiction for dispute resolution",
                "rationale": "Governing law clauses provide certainty about which legal system applies to disputes.",
                "category": "Legal Framework",
                "priority": "medium"
            })
        
        if risk_analysis["intellectual_property"]:
            recommendations.append({
                "recommendation": "Clarify intellectual property ownership and usage rights",
                "rationale": "IP-related terms were identified. Ensure clear ownership and usage rights are defined.",
                "category": "Intellectual Property",
                "priority": "high"
            })
        
        if risk_analysis["data_protection"]:
            recommendations.append({
                "recommendation": "Ensure compliance with data protection regulations (GDPR, CCPA, etc.)",
                "rationale": "Data protection considerations were found. Verify compliance with applicable privacy laws.",
                "category": "Compliance",
                "priority": "high"
            })
        
        # Generate general recommendations based on risk level
        high_risk_count = sum(1 for risk in risks if risk.get("level", "").lower() == "high")
        medium_risk_count = sum(1 for risk in risks if risk.get("level", "").lower() == "medium")
        
        if high_risk_count > 0:
            recommendations.append({
                "recommendation": "Schedule a legal review session to address high-priority risks",
                "rationale": f"{high_risk_count} high-risk issues were identified. Professional legal review is recommended.",
                "category": "Legal Review",
                "priority": "high"
            })
        
        if medium_risk_count > 2:
            recommendations.append({
                "recommendation": "Consider negotiating key terms before signing",
                "rationale": f"Multiple medium-risk issues ({medium_risk_count}) suggest the need for contract negotiation.",
                "category": "Negotiation",
                "priority": "medium"
            })
        
        # Add general best practices
        if not existing_recommendations:
            recommendations.append({
                "recommendation": "Review the contract with stakeholders before final approval",
                "rationale": "Multi-stakeholder review helps identify business-specific concerns and requirements.",
                "category": "Process",
                "priority": "low"
            })
            
            recommendations.append({
                "recommendation": "Document any verbal agreements or side arrangements",
                "rationale": "Ensure all important terms are captured in writing to avoid misunderstandings.",
                "category": "Documentation",
                "priority": "low"
            })
        
        return recommendations[:8]  # Limit to 8 recommendations to avoid overwhelming the user
    
    def get_history_data():
        """Get review history"""
        try:
            response = requests.get("http://localhost:7860/api/history")
            if response.status_code == 200:
                history = response.json().get("history", [])
                if not history:
                    return "No review history available."
                
                html = """
                <div class="result-section">
                    <h3>📋 Review History</h3>
                    <table style="width: 100%; border-collapse: collapse;">
                        <thead>
                            <tr style="background: #f8f9fa;">
                                <th style="padding: 0.5rem; border: 1px solid #dee2e6;">File</th>
                                <th style="padding: 0.5rem; border: 1px solid #dee2e6;">Date</th>
                                <th style="padding: 0.5rem; border: 1px solid #dee2e6;">Summary</th>
                                <th style="padding: 0.5rem; border: 1px solid #dee2e6;">Risks</th>
                            </tr>
                        </thead>
                        <tbody>
                """
                
                for item in history[-10:]:  # Show last 10
                    date = datetime.fromisoformat(item["timestamp"]).strftime("%Y-%m-%d %H:%M")
                    html += f"""
                    <tr>
                        <td style="padding: 0.5rem; border: 1px solid #dee2e6;">{item["filename"]}</td>
                        <td style="padding: 0.5rem; border: 1px solid #dee2e6;">{date}</td>
                        <td style="padding: 0.5rem; border: 1px solid #dee2e6;">{item["summary"][:100]}...</td>
                        <td style="padding: 0.5rem; border: 1px solid #dee2e6;">{item["risks_count"]}</td>
                    </tr>
                    """
                
                html += "</tbody></table></div>"
                return html
            else:
                return "Error loading history."
        except Exception as e:
            return f"Error: {str(e)}"
    
    # Create the modern UI
    with gr.Blocks(css=custom_css, title="Contract Review Assistant") as demo:
        gr.Markdown("""
        # 📋 Contract Review Assistant
        **AI-Powered Document Analysis & Risk Assessment**
        
        Upload your contracts and documents for comprehensive analysis, risk assessment, and actionable recommendations.
        """)
        
        with gr.Tabs():
            # Main Review Tab
            with gr.Tab("🔍 Document Review"):
                # Upload Section (collapsible)
                with gr.Group():
                    gr.Markdown("### 📤 Upload Document")
                    file_input = gr.File(
                        label="Upload PDF or DOCX file",
                        file_types=[".pdf", ".docx"],
                        file_count="single"
                    )
                    
                    with gr.Row():
                        process_btn = gr.Button("🚀 Analyze Document", variant="primary", size="lg")
                        upload_status = gr.Textbox(label="Status", visible=False)
                    
                    gr.Markdown("""
                    ### 📊 Supported Features
                    - **Document Analysis**: Extract key information and structure
                    - **Risk Assessment**: Identify potential risks and issues
                    - **Smart Recommendations**: Get actionable insights
                    - **Citation Tracking**: Find relevant sections and quotes
                    - **Export Options**: Download reports in multiple formats
                    """)
                
                # Results Section
                with gr.Group(visible=False) as results_group:
                    gr.Markdown("### 📋 Analysis Results")
                    
                    # Risk Overview
                    risk_overview = gr.HTML(label="Risk Overview")
                    
                    # Live status banner
                    status_html = gr.HTML(visible=False)
                    
                    # Main content in tabs
                    with gr.Tabs():
                        with gr.Tab("📄 Summary"):
                            summary_html = gr.HTML(label="Document Summary")
                        
                        with gr.Tab("🔑 Key Terms"):
                            key_terms_html = gr.HTML(label="Key Terms")
                        
                        with gr.Tab("⚠️ Risk Analysis"):
                            risks_html = gr.HTML(label="Risk Analysis")
                        
                        with gr.Tab("💡 Recommendations"):
                            recommendations_html = gr.HTML(label="Recommendations")
                        
                        with gr.Tab("💬 Comments"):
                            with gr.Row():
                                with gr.Column(scale=2):
                                    comments_display = gr.HTML(label="Comments", value="<div style='text-align: center; color: #6c757d; padding: 2rem;'>No comments yet. Add your first comment below!</div>")
                                with gr.Column(scale=1):
                                    gr.Markdown("### Add Comment")
                                    comment_text = gr.Textbox(
                                        label="Your comment",
                                        placeholder="Add your thoughts about this analysis...",
                                        lines=4
                                    )
                                    comment_section = gr.Dropdown(
                                        label="Section",
                                        choices=["Summary", "Risk Analysis", "Recommendations", "General"],
                                        value="General"
                                    )
                                    add_comment_btn = gr.Button("💬 Add Comment", variant="primary")
                        
                        with gr.Tab("🔧 Advanced"):
                            with gr.Row():
                                json_output = gr.Code(label="Raw JSON Results", language="json")
                                report_id_output = gr.Textbox(label="Report ID", visible=False)
                
                # Hidden state for showing results
                show_results = gr.State(False)
            
            # History Tab
            with gr.Tab("📚 Review History"):
                history_btn = gr.Button("🔄 Refresh History", variant="secondary")
                history_html = gr.HTML(label="History")
            
            # Settings Tab
            with gr.Tab("⚙️ Settings"):
                gr.Markdown("### 🔧 Application Settings")
                
                with gr.Row():
                    with gr.Column():
                        gr.Markdown(f"""
                        **Processing Options:**
                        - Document chunk size: 1200 characters
                        - Overlap: 120 characters
                        - Embedding model: legal-bert (768-dim)
                        - Chat model: llama3.2:7b
                        - Vector dimension: {EMBED_DIM}
                        """)
                    
                    with gr.Column():
                        gr.Markdown("""
                        **System Status:**
                        - Hub Gateway: Connected
                        - Vector Database: Active
                        - Cache: Enabled
                        - Reports: Stored locally
                        """)
                
                gr.Markdown("### 📁 Data Management")
                with gr.Row():
                    clear_history_btn = gr.Button("🗑️ Clear History", variant="stop")
                    export_data_btn = gr.Button("📤 Export Data", variant="secondary")
        
        # Event handlers
        def process_and_show_results(file):
            """Process document and show results with proper state management"""
            if file is None:
                return (None, "", "", "", "", "", "", gr.update(visible=False), "")
            
            # Process the document
            result = process_document(file)
            
            if result[0] is not None:  # Success
                json_data, summary, key_terms, risks, recommendations, report_id, show_results_flag, filename = result
                risk_overview_html = create_risk_overview(json.loads(json_data))
                
                return (
                    json_data,           # json_output
                    summary,            # summary_html
                    key_terms,          # key_terms_html
                    risks,              # risks_html
                    recommendations,    # recommendations_html
                    report_id,          # report_id_output
                    risk_overview_html, # risk_overview
                    gr.update(visible=True),
                    filename            # upload_status
                )
            else:
                # Error case
                return (None, "", "", "", "", "", "", gr.update(visible=False), result[1])
        
        # Helper functions to show/hide a live status banner
        def show_analyzing():
            spinner = """
            <div class=\"progress-container\" style=\"border:1px solid #dee2e6; border-radius:8px; padding:12px; margin-bottom:10px; background:#fff;\"> 
                <div style=\"display:flex; align-items:center; gap:8px; color:#495057;\">
                    <svg width=\"18\" height=\"18\" viewBox=\"0 0 24 24\" fill=\"none\" xmlns=\"http://www.w3.org/2000/svg\" style=\"animation: spin 1s linear infinite;\"><path d=\"M12 2a10 10 0 100 20 10 10 0 000-20zm0 3a7 7 0 110 14 7 7 0 010-14z\" fill=\"#0d6efd\"/></svg>
                    <span>Analyzing document… extracting text, chunking, embedding, and assessing risks.</span>
                </div>
                <div class=\"progress-bar\"><div class=\"progress-fill\" style=\"width:35%\"></div></div>
                <style>@keyframes spin{from{transform:rotate(0)}to{transform:rotate(360deg)}}</style>
            </div>
            """
            return gr.update(value=spinner, visible=True)
        
        def hide_status():
            return gr.update(visible=False)
        
        def show_results_immediately():
            """Show results section immediately when analysis starts"""
            return gr.update(visible=True)
        
        # Chain events: show results immediately -> show status -> run analysis -> hide status
        process_btn.click(
            fn=show_results_immediately,
            outputs=[results_group],
            show_progress=False
        ).then(
            fn=show_analyzing,
            outputs=[status_html],
            show_progress=False
        ).then(
            fn=process_and_show_results,
            inputs=[file_input],
            outputs=[json_output, summary_html, key_terms_html, risks_html, recommendations_html, report_id_output, risk_overview, results_group, upload_status]
        ).then(
            fn=hide_status,
            outputs=[status_html],
            show_progress=False
        )
        
        # Comment functionality
        def handle_add_comment(comment_text, section, report_id):
            if not report_id:
                return "No document loaded. Please analyze a document first.", gr.update()
            return add_comment(comment_text, section, report_id)
        
        add_comment_btn.click(
            fn=handle_add_comment,
            inputs=[comment_text, comment_section, report_id_output],
            outputs=[comment_text, comments_display]  # Clear comment text and update display
        )
        
        history_btn.click(
            fn=get_history_data,
            outputs=[history_html]
        )
        
        # Load history on tab switch
        demo.load(
            fn=get_history_data,
            outputs=[history_html]
        )
    
    # Enable Gradio task queue for better UX on long operations (use supported args)
    demo.queue(status_update_rate=1)
    return demo

if __name__ == "__main__":
    ui = create_modern_ui()
    # Mount Gradio Blocks on the FastAPI app root path
    app = gr.mount_gradio_app(app, ui, path="/")
    # Run a single uvicorn server to serve FastAPI (including mounted Gradio UI)
    uvicorn.run(app, host="0.0.0.0", port=APP_PORT)

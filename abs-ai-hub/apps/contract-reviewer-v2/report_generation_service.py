"""
Report Generation Service
Generate comprehensive reports in multiple formats (PDF, Word, JSON) with file-based storage
"""

import os
import json
import uuid
from pathlib import Path
from typing import List, Dict, Optional, Any, Union
from datetime import datetime
from enum import Enum
import logging
import asyncio
from dataclasses import dataclass
import tempfile

# Report generation libraries
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from file_based_storage_service import FileBasedStorageService, FileType, StorageConfig, FileMetadata

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class ReportFormat(Enum):
    """Report format enumeration"""
    PDF = "pdf"
    WORD = "docx"
    JSON = "json"
    HTML = "html"
    XML = "xml"


class ReportType(Enum):
    """Report type enumeration"""
    ANALYSIS_SUMMARY = "analysis_summary"
    DETAILED_ANALYSIS = "detailed_analysis"
    COMPARATIVE_ANALYSIS = "comparative_analysis"
    COMPLIANCE_REPORT = "compliance_report"
    EXECUTIVE_SUMMARY = "executive_summary"
    TECHNICAL_REPORT = "technical_report"


@dataclass
class ReportTemplate:
    """Report template structure"""
    template_id: str
    name: str
    description: str
    report_type: ReportType
    format: ReportFormat
    template_data: Dict[str, Any]
    created_at: datetime
    updated_at: datetime
    is_active: bool = True


@dataclass
class ReportRequest:
    """Report generation request"""
    report_id: str
    report_type: ReportType
    format: ReportFormat
    document_ids: List[str]
    analysis_ids: List[str]
    client_id: Optional[str] = None
    template_id: Optional[str] = None
    custom_data: Optional[Dict[str, Any]] = None
    include_charts: bool = True
    include_appendix: bool = True


class ReportGenerationService:
    """Service for generating comprehensive reports with file-based storage"""
    
    def __init__(self, storage_service: FileBasedStorageService):
        self.storage_service = storage_service
        self.templates: Dict[str, ReportTemplate] = {}
        self._load_default_templates()
    
    def _load_default_templates(self):
        """Load default report templates"""
        try:
            # Analysis Summary Template
            analysis_summary_template = ReportTemplate(
                template_id="analysis_summary_default",
                name="Analysis Summary Report",
                description="Standard analysis summary report",
                report_type=ReportType.ANALYSIS_SUMMARY,
                format=ReportFormat.PDF,
                template_data={
                    "sections": [
                        "executive_summary",
                        "key_findings",
                        "risk_assessment",
                        "recommendations",
                        "appendix"
                    ],
                    "styling": {
                        "font_family": "Helvetica",
                        "font_size": 12,
                        "line_spacing": 1.2,
                        "margins": {"top": 1, "bottom": 1, "left": 1, "right": 1}
                    }
                },
                created_at=datetime.now(),
                updated_at=datetime.now()
            )
            
            # Detailed Analysis Template
            detailed_analysis_template = ReportTemplate(
                template_id="detailed_analysis_default",
                name="Detailed Analysis Report",
                description="Comprehensive detailed analysis report",
                report_type=ReportType.DETAILED_ANALYSIS,
                format=ReportFormat.PDF,
                template_data={
                    "sections": [
                        "executive_summary",
                        "document_overview",
                        "detailed_analysis",
                        "clause_analysis",
                        "risk_assessment",
                        "recommendations",
                        "compliance_check",
                        "appendix"
                    ],
                    "styling": {
                        "font_family": "Helvetica",
                        "font_size": 11,
                        "line_spacing": 1.15,
                        "margins": {"top": 0.75, "bottom": 0.75, "left": 0.75, "right": 0.75}
                    }
                },
                created_at=datetime.now(),
                updated_at=datetime.now()
            )
            
            # Executive Summary Template
            executive_summary_template = ReportTemplate(
                template_id="executive_summary_default",
                name="Executive Summary Report",
                description="High-level executive summary report",
                report_type=ReportType.EXECUTIVE_SUMMARY,
                format=ReportFormat.PDF,
                template_data={
                    "sections": [
                        "executive_summary",
                        "key_risks",
                        "recommendations",
                        "next_steps"
                    ],
                    "styling": {
                        "font_family": "Helvetica",
                        "font_size": 14,
                        "line_spacing": 1.3,
                        "margins": {"top": 1, "bottom": 1, "left": 1, "right": 1}
                    }
                },
                created_at=datetime.now(),
                updated_at=datetime.now()
            )
            
            # Store templates
            self.templates[analysis_summary_template.template_id] = analysis_summary_template
            self.templates[detailed_analysis_template.template_id] = detailed_analysis_template
            self.templates[executive_summary_template.template_id] = executive_summary_template
            
            logger.info(f"✅ Loaded {len(self.templates)} default templates")
            
        except Exception as e:
            logger.error(f"❌ Error loading default templates: {e}")
    
    # ==================== REPORT GENERATION ====================
    
    async def generate_report(
        self,
        request: ReportRequest,
        analysis_data: Dict[str, Any],
        document_metadata: Optional[Dict[str, Any]] = None
    ) -> FileMetadata:
        """
        Generate a comprehensive report
        
        Args:
            request: Report generation request
            analysis_data: Analysis data to include in report
            document_metadata: Document metadata
            
        Returns:
            FileMetadata for the generated report
        """
        try:
            logger.info(f"Generating report: {request.report_id} ({request.format.value})")
            
            # Get template
            template = self._get_template(request.template_id, request.report_type, request.format)
            
            # Generate report content based on format
            if request.format == ReportFormat.PDF:
                report_content = await self._generate_pdf_report(
                    request, analysis_data, document_metadata, template
                )
            elif request.format == ReportFormat.WORD:
                report_content = await self._generate_word_report(
                    request, analysis_data, document_metadata, template
                )
            elif request.format == ReportFormat.JSON:
                report_content = await self._generate_json_report(
                    request, analysis_data, document_metadata, template
                )
            elif request.format == ReportFormat.HTML:
                report_content = await self._generate_html_report(
                    request, analysis_data, document_metadata, template
                )
            else:
                raise ValueError(f"Unsupported report format: {request.format}")
            
            # Store report file
            filename = f"{request.report_type.value}_{request.report_id}.{request.format.value}"
            
            file_metadata = await self.storage_service.store_file(
                file_data=report_content,
                file_type=FileType.REPORT_PDF if request.format == ReportFormat.PDF else FileType.REPORT_WORD,
                original_filename=filename,
                client_id=request.client_id,
                document_id=request.document_ids[0] if request.document_ids else None,
                analysis_id=request.analysis_ids[0] if request.analysis_ids else None,
                metadata={
                    "report_type": request.report_type.value,
                    "report_format": request.format.value,
                    "template_id": template.template_id,
                    "generated_at": datetime.now().isoformat(),
                    "document_count": len(request.document_ids),
                    "analysis_count": len(request.analysis_ids),
                    "include_charts": request.include_charts,
                    "include_appendix": request.include_appendix
                }
            )
            
            logger.info(f"✅ Report generated: {file_metadata.file_id}")
            return file_metadata
            
        except Exception as e:
            logger.error(f"❌ Error generating report: {e}")
            raise
    
    def _get_template(
        self,
        template_id: Optional[str],
        report_type: ReportType,
        format: ReportFormat
    ) -> ReportTemplate:
        """Get report template"""
        if template_id and template_id in self.templates:
            return self.templates[template_id]
        
        # Find default template for report type and format
        for template in self.templates.values():
            if template.report_type == report_type and template.format == format:
                return template
        
        # Fallback to first available template
        if self.templates:
            return list(self.templates.values())[0]
        
        raise ValueError("No templates available")
    
    # ==================== PDF REPORT GENERATION ====================
    
    async def _generate_pdf_report(
        self,
        request: ReportRequest,
        analysis_data: Dict[str, Any],
        document_metadata: Optional[Dict[str, Any]],
        template: ReportTemplate
    ) -> bytes:
        """Generate PDF report using ReportLab"""
        if not REPORTLAB_AVAILABLE:
            raise ImportError("ReportLab not available. Install with: pip install reportlab")
        
        try:
            # Create temporary file
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
                temp_path = temp_file.name
            
            try:
                # Create PDF document
                doc = SimpleDocTemplate(
                    temp_path,
                    pagesize=A4,
                    rightMargin=72,
                    leftMargin=72,
                    topMargin=72,
                    bottomMargin=18
                )
                
                # Get styles
                styles = getSampleStyleSheet()
                
                # Create custom styles
                title_style = ParagraphStyle(
                    'CustomTitle',
                    parent=styles['Heading1'],
                    fontSize=18,
                    spaceAfter=30,
                    alignment=1  # Center alignment
                )
                
                heading_style = ParagraphStyle(
                    'CustomHeading',
                    parent=styles['Heading2'],
                    fontSize=14,
                    spaceAfter=12,
                    spaceBefore=20
                )
                
                body_style = ParagraphStyle(
                    'CustomBody',
                    parent=styles['Normal'],
                    fontSize=11,
                    spaceAfter=6,
                    leading=14
                )
                
                # Build report content
                story = []
                
                # Title
                story.append(Paragraph(f"{template.name}", title_style))
                story.append(Spacer(1, 20))
                
                # Report metadata
                metadata_table_data = [
                    ["Report ID:", request.report_id],
                    ["Generated:", datetime.now().strftime("%Y-%m-%d %H:%M:%S")],
                    ["Report Type:", request.report_type.value.replace('_', ' ').title()],
                    ["Format:", request.format.value.upper()]
                ]
                
                if request.client_id:
                    metadata_table_data.append(["Client:", request.client_id])
                
                if document_metadata:
                    metadata_table_data.extend([
                        ["Document:", document_metadata.get("original_filename", "Unknown")],
                        ["File Size:", f"{document_metadata.get('file_size', 0) / 1024:.1f} KB"],
                        ["Upload Date:", document_metadata.get("upload_timestamp", "Unknown")]
                    ])
                
                metadata_table = Table(metadata_table_data, colWidths=[2*inch, 4*inch])
                metadata_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
                    ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                    ('BACKGROUND', (1, 0), (1, -1), colors.beige),
                ]))
                
                story.append(metadata_table)
                story.append(Spacer(1, 20))
                
                # Executive Summary
                story.append(Paragraph("Executive Summary", heading_style))
                summary_text = analysis_data.get("summary", {}).get("summary", "No summary available.")
                story.append(Paragraph(summary_text, body_style))
                story.append(Spacer(1, 12))
                
                # Key Findings
                story.append(Paragraph("Key Findings", heading_style))
                key_points = analysis_data.get("summary", {}).get("key_points", [])
                if key_points:
                    for point in key_points:
                        story.append(Paragraph(f"• {point}", body_style))
                else:
                    story.append(Paragraph("No key findings identified.", body_style))
                story.append(Spacer(1, 12))
                
                # Risk Assessment
                story.append(Paragraph("Risk Assessment", heading_style))
                risks = analysis_data.get("risks", [])
                if risks:
                    risk_table_data = [["Risk Level", "Description"]]
                    for risk in risks:
                        risk_table_data.append([
                            risk.get("level", "Unknown").upper(),
                            risk.get("description", "No description")
                        ])
                    
                    risk_table = Table(risk_table_data, colWidths=[1.5*inch, 4.5*inch])
                    risk_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 10),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    
                    story.append(risk_table)
                else:
                    story.append(Paragraph("No risks identified.", body_style))
                story.append(Spacer(1, 12))
                
                # Recommendations
                story.append(Paragraph("Recommendations", heading_style))
                recommendations = analysis_data.get("recommendations", [])
                if recommendations:
                    for i, rec in enumerate(recommendations, 1):
                        story.append(Paragraph(f"{i}. {rec}", body_style))
                else:
                    story.append(Paragraph("No recommendations provided.", body_style))
                story.append(Spacer(1, 12))
                
                # Citations
                citations = analysis_data.get("citations", [])
                if citations:
                    story.append(Paragraph("References", heading_style))
                    for citation in citations:
                        story.append(Paragraph(f"• {citation}", body_style))
                
                # Build PDF
                doc.build(story)
                
                # Read generated PDF
                with open(temp_path, 'rb') as f:
                    pdf_content = f.read()
                
                return pdf_content
                
            finally:
                # Clean up temp file
                os.unlink(temp_path)
                
        except Exception as e:
            logger.error(f"❌ Error generating PDF report: {e}")
            raise
    
    # ==================== WORD REPORT GENERATION ====================
    
    async def _generate_word_report(
        self,
        request: ReportRequest,
        analysis_data: Dict[str, Any],
        document_metadata: Optional[Dict[str, Any]],
        template: ReportTemplate
    ) -> bytes:
        """Generate Word report using python-docx"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available. Install with: pip install python-docx")
        
        try:
            # Create Word document
            doc = DocxDocument()
            
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Title
            title = doc.add_heading(template.name, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Report metadata
            doc.add_heading('Report Information', level=1)
            
            metadata_para = doc.add_paragraph()
            metadata_para.add_run(f"Report ID: ").bold = True
            metadata_para.add_run(f"{request.report_id}\n")
            metadata_para.add_run(f"Generated: ").bold = True
            metadata_para.add_run(f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            metadata_para.add_run(f"Report Type: ").bold = True
            metadata_para.add_run(f"{request.report_type.value.replace('_', ' ').title()}\n")
            metadata_para.add_run(f"Format: ").bold = True
            metadata_para.add_run(f"{request.format.value.upper()}\n")
            
            if request.client_id:
                metadata_para.add_run(f"Client: ").bold = True
                metadata_para.add_run(f"{request.client_id}\n")
            
            if document_metadata:
                metadata_para.add_run(f"Document: ").bold = True
                metadata_para.add_run(f"{document_metadata.get('original_filename', 'Unknown')}\n")
                metadata_para.add_run(f"File Size: ").bold = True
                metadata_para.add_run(f"{document_metadata.get('file_size', 0) / 1024:.1f} KB\n")
                metadata_para.add_run(f"Upload Date: ").bold = True
                metadata_para.add_run(f"{document_metadata.get('upload_timestamp', 'Unknown')}\n")
            
            # Executive Summary
            doc.add_heading('Executive Summary', level=1)
            summary_text = analysis_data.get("summary", {}).get("summary", "No summary available.")
            doc.add_paragraph(summary_text)
            
            # Key Findings
            doc.add_heading('Key Findings', level=1)
            key_points = analysis_data.get("summary", {}).get("key_points", [])
            if key_points:
                for point in key_points:
                    p = doc.add_paragraph(point, style='List Bullet')
            else:
                doc.add_paragraph("No key findings identified.")
            
            # Risk Assessment
            doc.add_heading('Risk Assessment', level=1)
            risks = analysis_data.get("risks", [])
            if risks:
                # Create risk table
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                
                # Header row
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Risk Level'
                hdr_cells[1].text = 'Description'
                
                # Add risk data
                for risk in risks:
                    row_cells = table.add_row().cells
                    row_cells[0].text = risk.get("level", "Unknown").upper()
                    row_cells[1].text = risk.get("description", "No description")
            else:
                doc.add_paragraph("No risks identified.")
            
            # Recommendations
            doc.add_heading('Recommendations', level=1)
            recommendations = analysis_data.get("recommendations", [])
            if recommendations:
                for i, rec in enumerate(recommendations, 1):
                    doc.add_paragraph(f"{i}. {rec}")
            else:
                doc.add_paragraph("No recommendations provided.")
            
            # Citations
            citations = analysis_data.get("citations", [])
            if citations:
                doc.add_heading('References', level=1)
                for citation in citations:
                    doc.add_paragraph(citation, style='List Bullet')
            
            # Save to bytes
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_path = temp_file.name
            
            try:
                doc.save(temp_path)
                
                with open(temp_path, 'rb') as f:
                    docx_content = f.read()
                
                return docx_content
                
            finally:
                os.unlink(temp_path)
                
        except Exception as e:
            logger.error(f"❌ Error generating Word report: {e}")
            raise
    
    # ==================== JSON REPORT GENERATION ====================
    
    async def _generate_json_report(
        self,
        request: ReportRequest,
        analysis_data: Dict[str, Any],
        document_metadata: Optional[Dict[str, Any]],
        template: ReportTemplate
    ) -> bytes:
        """Generate JSON report"""
        try:
            report_data = {
                "report_metadata": {
                    "report_id": request.report_id,
                    "report_type": request.report_type.value,
                    "format": request.format.value,
                    "template_id": template.template_id,
                    "generated_at": datetime.now().isoformat(),
                    "client_id": request.client_id,
                    "document_count": len(request.document_ids),
                    "analysis_count": len(request.analysis_ids)
                },
                "document_metadata": document_metadata,
                "analysis_data": analysis_data,
                "template_info": {
                    "template_id": template.template_id,
                    "template_name": template.name,
                    "template_description": template.description
                },
                "generation_info": {
                    "include_charts": request.include_charts,
                    "include_appendix": request.include_appendix,
                    "custom_data": request.custom_data
                }
            }
            
            json_content = json.dumps(report_data, indent=2, default=str)
            return json_content.encode('utf-8')
            
        except Exception as e:
            logger.error(f"❌ Error generating JSON report: {e}")
            raise
    
    # ==================== HTML REPORT GENERATION ====================
    
    async def _generate_html_report(
        self,
        request: ReportRequest,
        analysis_data: Dict[str, Any],
        document_metadata: Optional[Dict[str, Any]],
        template: ReportTemplate
    ) -> bytes:
        """Generate HTML report"""
        try:
            html_content = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{template.name}</title>
    <style>
        body {{
            font-family: Arial, sans-serif;
            line-height: 1.6;
            margin: 0;
            padding: 20px;
            background-color: #f4f4f4;
        }}
        .container {{
            max-width: 800px;
            margin: 0 auto;
            background-color: white;
            padding: 30px;
            border-radius: 10px;
            box-shadow: 0 0 10px rgba(0,0,0,0.1);
        }}
        h1 {{
            color: #333;
            text-align: center;
            border-bottom: 2px solid #007acc;
            padding-bottom: 10px;
        }}
        h2 {{
            color: #007acc;
            margin-top: 30px;
        }}
        .metadata {{
            background-color: #f8f9fa;
            padding: 15px;
            border-radius: 5px;
            margin-bottom: 20px;
        }}
        .risk-table {{
            width: 100%;
            border-collapse: collapse;
            margin: 15px 0;
        }}
        .risk-table th, .risk-table td {{
            border: 1px solid #ddd;
            padding: 8px;
            text-align: left;
        }}
        .risk-table th {{
            background-color: #007acc;
            color: white;
        }}
        .risk-low {{ background-color: #d4edda; }}
        .risk-medium {{ background-color: #fff3cd; }}
        .risk-high {{ background-color: #f8d7da; }}
        ul {{
            padding-left: 20px;
        }}
        .footer {{
            margin-top: 40px;
            padding-top: 20px;
            border-top: 1px solid #ddd;
            text-align: center;
            color: #666;
        }}
    </style>
</head>
<body>
    <div class="container">
        <h1>{template.name}</h1>
        
        <div class="metadata">
            <h2>Report Information</h2>
            <p><strong>Report ID:</strong> {request.report_id}</p>
            <p><strong>Generated:</strong> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
            <p><strong>Report Type:</strong> {request.report_type.value.replace('_', ' ').title()}</p>
            <p><strong>Format:</strong> {request.format.value.upper()}</p>
"""
            
            if request.client_id:
                html_content += f'            <p><strong>Client:</strong> {request.client_id}</p>\n'
            
            if document_metadata:
                html_content += f"""
            <p><strong>Document:</strong> {document_metadata.get('original_filename', 'Unknown')}</p>
            <p><strong>File Size:</strong> {document_metadata.get('file_size', 0) / 1024:.1f} KB</p>
            <p><strong>Upload Date:</strong> {document_metadata.get('upload_timestamp', 'Unknown')}</p>
"""
            
            html_content += """
        </div>
        
        <h2>Executive Summary</h2>
        <p>"""
            
            summary_text = analysis_data.get("summary", {}).get("summary", "No summary available.")
            html_content += summary_text.replace('\n', '<br>')
            
            html_content += """
        </p>
        
        <h2>Key Findings</h2>
        <ul>
"""
            
            key_points = analysis_data.get("summary", {}).get("key_points", [])
            if key_points:
                for point in key_points:
                    html_content += f"            <li>{point}</li>\n"
            else:
                html_content += "            <li>No key findings identified.</li>\n"
            
            html_content += """
        </ul>
        
        <h2>Risk Assessment</h2>
"""
            
            risks = analysis_data.get("risks", [])
            if risks:
                html_content += """
        <table class="risk-table">
            <thead>
                <tr>
                    <th>Risk Level</th>
                    <th>Description</th>
                </tr>
            </thead>
            <tbody>
"""
                for risk in risks:
                    risk_level = risk.get("level", "unknown").lower()
                    risk_class = f"risk-{risk_level}"
                    html_content += f"""
                <tr class="{risk_class}">
                    <td>{risk.get('level', 'Unknown').upper()}</td>
                    <td>{risk.get('description', 'No description')}</td>
                </tr>
"""
                html_content += """
            </tbody>
        </table>
"""
            else:
                html_content += "        <p>No risks identified.</p>\n"
            
            html_content += """
        
        <h2>Recommendations</h2>
        <ol>
"""
            
            recommendations = analysis_data.get("recommendations", [])
            if recommendations:
                for rec in recommendations:
                    html_content += f"            <li>{rec}</li>\n"
            else:
                html_content += "            <li>No recommendations provided.</li>\n"
            
            html_content += """
        </ol>
"""
            
            citations = analysis_data.get("citations", [])
            if citations:
                html_content += """
        
        <h2>References</h2>
        <ul>
"""
                for citation in citations:
                    html_content += f"            <li>{citation}</li>\n"
                html_content += "        </ul>\n"
            
            html_content += f"""
        
        <div class="footer">
            <p>Generated by Contract Reviewer v2 - Enhanced with Vector Search</p>
            <p>Report ID: {request.report_id} | Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
        </div>
    </div>
</body>
</html>
"""
            
            return html_content.encode('utf-8')
            
        except Exception as e:
            logger.error(f"❌ Error generating HTML report: {e}")
            raise
    
    # ==================== TEMPLATE MANAGEMENT ====================
    
    async def create_template(
        self,
        template_id: str,
        name: str,
        description: str,
        report_type: ReportType,
        format: ReportFormat,
        template_data: Dict[str, Any]
    ) -> ReportTemplate:
        """Create a new report template"""
        try:
            template = ReportTemplate(
                template_id=template_id,
                name=name,
                description=description,
                report_type=report_type,
                format=format,
                template_data=template_data,
                created_at=datetime.now(),
                updated_at=datetime.now()
            )
            
            self.templates[template_id] = template
            
            # Store template in file system
            await self.storage_service.store_file(
                file_data=template_data,
                file_type=FileType.TEMPLATE,
                original_filename=f"template_{template_id}.json",
                metadata={
                    "template_id": template_id,
                    "name": name,
                    "description": description,
                    "report_type": report_type.value,
                    "format": format.value,
                    "created_at": template.created_at.isoformat(),
                    "updated_at": template.updated_at.isoformat()
                }
            )
            
            logger.info(f"✅ Template created: {template_id}")
            return template
            
        except Exception as e:
            logger.error(f"❌ Error creating template: {e}")
            raise
    
    async def get_template(self, template_id: str) -> Optional[ReportTemplate]:
        """Get template by ID"""
        return self.templates.get(template_id)
    
    async def list_templates(self, report_type: Optional[ReportType] = None) -> List[ReportTemplate]:
        """List available templates"""
        templates = list(self.templates.values())
        
        if report_type:
            templates = [t for t in templates if t.report_type == report_type]
        
        return templates
    
    # ==================== UTILITY METHODS ====================
    
    async def get_report_stats(self) -> Dict[str, Any]:
        """Get report generation statistics"""
        try:
            stats = await self.storage_service.get_storage_stats()
            
            # Filter for report files
            report_files = 0
            report_size = 0
            
            for file_type, count in stats.get("files_by_type", {}).items():
                if file_type in ["report_pdf", "report_word", "report_json"]:
                    report_files += count
            
            # This is a simplified calculation - in practice, you'd need to iterate through files
            report_size = stats.get("total_size_bytes", 0) * 0.1  # Estimate 10% are reports
            
            return {
                "total_templates": len(self.templates),
                "templates_by_type": {
                    template.report_type.value: len([t for t in self.templates.values() if t.report_type == template.report_type])
                    for template in self.templates.values()
                },
                "templates_by_format": {
                    template.format.value: len([t for t in self.templates.values() if t.format == template.format])
                    for template in self.templates.values()
                },
                "total_reports": report_files,
                "total_report_size_mb": report_size / (1024 * 1024),
                "available_formats": [fmt.value for fmt in ReportFormat],
                "available_types": [rtype.value for rtype in ReportType]
            }
            
        except Exception as e:
            logger.error(f"❌ Error getting report stats: {e}")
            raise


# ==================== EXAMPLE USAGE ====================

async def example_usage():
    """Example of how to use the ReportGenerationService"""
    
    # Initialize services
    config = StorageConfig(base_path="/tmp/report_storage_example")
    storage_service = FileBasedStorageService(config)
    report_service = ReportGenerationService(storage_service)
    
    try:
        # Sample analysis data
        analysis_data = {
            "summary": {
                "summary": "This is a comprehensive analysis of the confidentiality agreement.",
                "key_points": [
                    "Standard confidentiality period of 2 years",
                    "Clear definition of confidential information",
                    "Appropriate remedies for breach"
                ]
            },
            "risks": [
                {"level": "low", "description": "Standard confidentiality clause"},
                {"level": "medium", "description": "Consider adding return of materials clause"}
            ],
            "recommendations": [
                "Review confidentiality period",
                "Add return of materials clause",
                "Verify governing law jurisdiction"
            ],
            "citations": [
                "Section 2.1: Confidentiality obligations",
                "Section 4.2: Term and termination"
            ]
        }
        
        # Sample document metadata
        document_metadata = {
            "original_filename": "confidentiality_agreement.pdf",
            "file_size": 102400,
            "upload_timestamp": "2024-01-15T10:30:00"
        }
        
        # Generate PDF report
        pdf_request = ReportRequest(
            report_id="report-001",
            report_type=ReportType.ANALYSIS_SUMMARY,
            format=ReportFormat.PDF,
            document_ids=["doc-001"],
            analysis_ids=["analysis-001"],
            client_id="ACME_Corp"
        )
        
        pdf_metadata = await report_service.generate_report(
            request=pdf_request,
            analysis_data=analysis_data,
            document_metadata=document_metadata
        )
        
        print(f"PDF report generated: {pdf_metadata.file_id}")
        
        # Generate Word report
        word_request = ReportRequest(
            report_id="report-002",
            report_type=ReportType.DETAILED_ANALYSIS,
            format=ReportFormat.WORD,
            document_ids=["doc-001"],
            analysis_ids=["analysis-001"],
            client_id="ACME_Corp"
        )
        
        word_metadata = await report_service.generate_report(
            request=word_request,
            analysis_data=analysis_data,
            document_metadata=document_metadata
        )
        
        print(f"Word report generated: {word_metadata.file_id}")
        
        # Generate JSON report
        json_request = ReportRequest(
            report_id="report-003",
            report_type=ReportType.TECHNICAL_REPORT,
            format=ReportFormat.JSON,
            document_ids=["doc-001"],
            analysis_ids=["analysis-001"],
            client_id="ACME_Corp"
        )
        
        json_metadata = await report_service.generate_report(
            request=json_request,
            analysis_data=analysis_data,
            document_metadata=document_metadata
        )
        
        print(f"JSON report generated: {json_metadata.file_id}")
        
        # Get report statistics
        stats = await report_service.get_report_stats()
        print(f"Report statistics: {stats}")
        
    finally:
        # Cleanup
        import shutil
        shutil.rmtree(config.base_path, ignore_errors=True)


if __name__ == "__main__":
    asyncio.run(example_usage())
